from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
DOCX_PATH = OUTPUT_DIR / "fragpunk-marketing-case-import-test.docx"
IMAGE_PATH = TMP_DIR / "fragpunk-import-test-key-art.png"


def make_test_image() -> None:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1400, 780), "#171b22")
    draw = ImageDraw.Draw(image)

    for index in range(0, 1400, 18):
        color = (35 + index % 90, 78, 96 + index % 120)
        draw.line((index, 0, 1400 - index // 3, 780), fill=color, width=2)

    draw.rectangle((90, 90, 1310, 690), outline="#42d6b5", width=6)
    draw.rectangle((130, 130, 1270, 650), fill="#0d1118")
    draw.polygon([(200, 580), (420, 260), (660, 580)], fill="#ff4b5c")
    draw.polygon([(520, 580), (780, 180), (1090, 580)], fill="#42d6b5")
    draw.ellipse((1040, 160, 1200, 320), fill="#f7c948")

    try:
        font_large = ImageFont.truetype("Arial Bold.ttf", 78)
        font_small = ImageFont.truetype("Arial.ttf", 34)
    except OSError:
        font_large = ImageFont.load_default()
        font_small = ImageFont.load_default()

    draw.text((160, 150), "FRAGPUNK", fill="#ffffff", font=font_large)
    draw.text((164, 250), "Import test key visual", fill="#c9d6df", font=font_small)
    draw.text((164, 312), "Image block with caption", fill="#c9d6df", font=font_small)
    image.save(IMAGE_PATH)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    text_node = OxmlElement("w:t")
    text_node.text = text

    run.append(run_properties)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        heading = document.styles[name]
        heading.font.name = "Arial"
        heading.font.color.rgb = RGBColor(31, 31, 27)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    make_test_image()

    document = Document()
    set_normal_style(document)

    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    document.add_heading("FragPunk 海外创意营销案例导入测试稿", level=1)
    intro = document.add_paragraph()
    intro.add_run("摘要：").bold = True
    intro.add_run(
        "这是一篇用于测试后台 Word 导入的完整稿件，覆盖标题层级、正文、引用、列表、图片、图片链接、YouTube、TikTok 和 MP4 视频链接。"
    )

    document.add_heading("一、核心案例：用规则反转制造短视频传播点", level=2)
    document.add_paragraph(
        "FragPunk 的传播重点不是简单告诉海外玩家这是一款 5v5 射击游戏，而是让玩家快速看懂 Shard Card 如何改变每一局的对抗节奏。"
        "测试稿中这一段应导入为普通正文段落，并在公开页保持舒适的阅读宽度。"
    )
    document.add_paragraph(
        "好的创意营销不是替玩家说话，而是把玩家已经在讨论的高光瞬间放大到更容易被看见。",
        style="Intense Quote",
    )

    document.add_heading("二、图片块：内嵌图片和图片说明", level=2)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(IMAGE_PATH), width=Inches(5.9))
    document.add_paragraph("图注：这张图是 Word 内嵌图片，导入后应保存为本地上传图片并显示在文章中。")

    document.add_heading("三、图片 URL 独立成行", level=2)
    document.add_paragraph("图片：https://images.unsplash.com/photo-1550745165-9bc0b252726f?auto=format&fit=crop&w=1600&q=80")
    document.add_paragraph("图注：这是一条图片 URL，粘贴解析时应识别为图片模块。")

    document.add_heading("四、视频链接：YouTube 与 TikTok", level=2)
    document.add_paragraph("YouTube：https://www.youtube.com/watch?v=dQw4w9WgXcQ")
    document.add_paragraph("视频说明：这条 YouTube 链接应在公开页转为 iframe 预览。")
    document.add_paragraph("TikTok：https://www.tiktok.com/@tiktok/video/7255302408174423339")
    document.add_paragraph("视频说明：这条 TikTok 链接应在公开页转为 TikTok iframe 预览；如果外部平台限制加载，至少应保留链接。")

    document.add_heading("五、直接视频文件和普通链接", level=2)
    document.add_paragraph("MP4：https://interactive-examples.mdn.mozilla.net/media/cc0-videos/flower.mp4")
    link_paragraph = document.add_paragraph("普通参考链接：")
    add_hyperlink(link_paragraph, "Steam 商店页示例", "https://store.steampowered.com/")

    document.add_heading("六、列表与分割线", level=2)
    document.add_paragraph("测试要点：", style=None)
    for item in ["标题层级能被保留", "正文段落能被拆分", "图片和视频链接能被识别", "引用和列表能保持基本层级"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph("---")

    document.add_heading("七、结尾判断", level=2)
    document.add_paragraph(
        "如果这篇 Word 上传后，公开页能看到标题、正文、图片、视频预览和清晰的段落层级，就说明第一版导入链路已经足够支撑内部分享 MVP。"
    )

    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
